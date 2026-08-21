"""
Builds a .docx for a tender-required document we don't have a hand-authored
template for, by reproducing the tender's own text for that document
verbatim -- under Airox's letterhead/footer -- instead of asking an LLM to
invent content. Many bundled tender documents (Inspection & Test Plans,
Quality Assurance Plan forms, an Approved Makes list) are already complete,
ready-to-sign documents inside the tender PDF; the bidder's job is to
countersign them, not rewrite them. This keeps that guarantee: nothing in
the body text comes from anywhere but the tender itself.
"""
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import branding
from .branding import docx_safe as _docx_safe

FONT_NAME = "Times New Roman"
FONT_SIZE = 11

# Tender ITP/QAP forms routinely pack many side-by-side columns into what
# was a landscape-printed page (MECON's QAP forms run 16-22 columns once
# pdfplumber splits out every merged-cell fragment). Dividing a 6" portrait
# page that many ways leaves each column under a quarter-inch -- too narrow
# for even one 11pt character, so every word wraps onto its own line and the
# table reads as scrambled vertical letters. Past this many columns the page
# is switched to landscape and the table font shrunk to keep cells legible.
WIDE_TABLE_COL_THRESHOLD = 6


def slugify(name: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")
    return (slug or "document") + ".docx"


def _set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    """Word's default cell padding (~0.08"-0.1" each side) eats a big chunk
    of an already-narrow column; tightening it to ~0.03"-0.04" (units are
    twentieths of a point) reclaims usable width for the actual text."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _render_table(doc, table_data: list[list[str]], usable_width_in: float) -> None:
    """Render a pdfplumber-extracted table as a real Word table. Reproducing
    a tender page with side-by-side columns as flat paragraph text scrambles
    it into unreadable word-soup (no column boundaries survive); an actual
    table keeps each cell where it belongs."""
    if not table_data or not table_data[0]:
        return
    n_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"

    # Distribute the page's usable width evenly, then pick a font size that
    # actually fits that column width instead of leaving the 11pt body font
    # in a column too narrow for it (the root cause of the letter-per-line
    # rendering on wide QAP/ITP tables).
    col_width_in = usable_width_in / n_cols
    font_pt = max(6, min(FONT_SIZE, round(col_width_in * 11)))
    col_width = Inches(col_width_in)
    for col in table.columns:
        col.width = col_width

    for row in table_data:
        cells = table.add_row().cells
        for i in range(n_cols):
            value = row[i] if i < len(row) else ""
            cell = cells[i]
            cell.width = col_width
            _set_cell_margins(cell)
            cell.text = ""
            run = cell.paragraphs[0].add_run(_docx_safe(value.strip()))
            run.font.size = Pt(font_pt)
            run.font.name = FONT_NAME
    doc.add_paragraph()


def build_verbatim_document(
    title: str,
    page_entries: list,
    output_path: str,
    company: dict,
    signing_date: str,
    signing_place: str,
) -> None:
    """page_entries: list of pages this document spans in the source tender
    PDF, in order. Each entry is either a plain (page_number, text) tuple
    (legacy form, always rendered as flat paragraphs) or a dict
    {"page": n, "text": str, "tables": list[list[list[str]]]} as returned by
    extraction.parse_pdf_page_range_detailed -- pages with genuine tables are
    rendered as real Word tables instead of flattened paragraph text."""
    max_cols = 0
    for entry in page_entries:
        if isinstance(entry, dict):
            for table_data in entry.get("tables") or []:
                if table_data:
                    max_cols = max(max_cols, max(len(row) for row in table_data))

    doc = Document()
    branding.set_default_font(doc, FONT_NAME, FONT_SIZE)
    if max_cols > WIDE_TABLE_COL_THRESHOLD:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    branding.add_letterhead(doc)
    branding.add_footer(doc, company)
    branding.add_title(doc, _docx_safe(title))

    usable_width_in = (
        doc.sections[0].page_width.inches
        - doc.sections[0].left_margin.inches
        - doc.sections[0].right_margin.inches
    )

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Reproduced verbatim from the tender document (no template was available "
        "for this item -- review against the source before signing)."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph()

    for entry in page_entries:
        if isinstance(entry, dict):
            page_no, text, tables = entry["page"], entry["text"], entry.get("tables") or []
        else:
            page_no, text, tables = entry[0], entry[1], []

        caption = doc.add_paragraph()
        caption_run = caption.add_run(f"[Tender page {page_no}]")
        caption_run.italic = True
        caption_run.font.size = Pt(8)
        caption_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # A page whose only detected "table" is a single-row box is really
        # just the page's title banner caught by pdfplumber's border
        # detection, not an actual data table -- narrative/clause pages
        # (e.g. numbered SCOPE/METHODOLOGY paragraphs) commonly get exactly
        # that. Treating it as "this page is a table page" would skip the
        # flat-text branch below and silently drop the entire page's real
        # content, leaving only the banner. A genuine checklist/ITP table
        # always has more than one row, so that's the signal used here.
        has_real_table = any(len(t) > 1 for t in tables)
        if has_real_table:
            # The page's tables already carry essentially all of its
            # meaningful content (checklist rows, form fields, headers);
            # also flattening the same text as paragraphs would duplicate it
            # and reintroduce the scrambled-column problem this exists to
            # avoid, so skip the flat-text rendering for these pages --
            # unless the table is nowhere near the whole page (see below).
            for table_data in tables:
                _render_table(doc, table_data, usable_width_in)

        # Some pages mix a small real table with substantial narrative
        # clauses that live entirely outside it (a numbered-methodology
        # page with a short "facility list" table dropped in the middle,
        # say) -- there the table is nowhere near the whole page, and
        # skipping flat text would silently drop that narrative. Measured
        # on real MRPL ITP/procedure pages: pages where the table
        # genuinely *is* the page have extract_text() length within
        # ~1.2-1.6x the table's own character count (extract_text() just
        # re-flattens the same cells); pages with real narrative alongside
        # a small table run 4.5-7x -- so a wide-margin ratio of 2.5 tells
        # the two apart without re-triggering on table-only pages.
        table_chars = sum(len(cell or "") for t in tables for row in t for cell in row)
        needs_flat_text = not has_real_table or len(text) > table_chars * 2.5
        if needs_flat_text:
            for line in text.splitlines():
                line = _docx_safe(line.strip())
                if not line:
                    continue
                doc.add_paragraph(line)
        doc.add_paragraph()

    branding.add_signoff_tail(doc, company, signing_place, signing_date)
    branding.ensure_unique_drawing_ids(doc)
    doc.save(output_path)


def build_ai_drafted_document(
    title: str,
    body_text: str,
    output_path: str,
    company: dict,
    signing_date: str,
    signing_place: str,
) -> None:
    """Builds a .docx for a tender-required document that has no
    hand-authored template AND no bundled text anywhere in the tender to
    reproduce (see extraction.compose_document_text -- this is only used
    for the has_own_content: false case; anything the tender itself
    provides text for still goes through build_verbatim_document above,
    never rewritten). The body text comes from the AI, but letterhead,
    footer, and signoff are the exact same shared Airox branding every
    generated document gets -- never anything AI-authored."""
    doc = Document()
    branding.set_default_font(doc, FONT_NAME, FONT_SIZE)
    branding.add_letterhead(doc)
    branding.add_footer(doc, company)
    branding.add_title(doc, _docx_safe(title))

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Drafted by Airox's AI assistant -- no ready-made format for this "
        "document exists anywhere in the tender. Review the content "
        "carefully, and fill in any [FILL: ...] items by hand, before "
        "signing or submission."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph()

    for line in body_text.splitlines():
        line = _docx_safe(line.strip())
        if not line:
            continue
        doc.add_paragraph(line)
    doc.add_paragraph()

    branding.add_signoff_tail(doc, company, signing_place, signing_date)
    branding.ensure_unique_drawing_ids(doc)
    doc.save(output_path)
