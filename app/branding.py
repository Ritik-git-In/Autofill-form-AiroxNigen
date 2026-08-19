"""
Airox's letterhead/footer/signoff, shared by anything that builds a .docx
at runtime (currently: the verbatim-document generator in generator.py).
build_templates.py has its own copies of this same logic -- that script is
a one-time reference tool, not part of the running app, so it's left alone
rather than refactored to share this module.

Signature and company stamp images are optional: if SIGNATURE_PATH /
STAMP_PATH exist on disk, the signoff block embeds them; otherwise it
leaves the same blank space the hand-typed templates always have for a
wet-ink signature, so this keeps working even if either file is removed.
"""
import os
import re

from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.dirname(__file__))
LOGO_PATH = os.path.join(HERE, "static", "airox_logo.jpg")
SIGNATURE_PATH = os.path.join(HERE, "static", "Signature Airox Nigen.jpg")
STAMP_PATH = os.path.join(HERE, "static", "Stamp Airox Nigen.jpg")

# Shared last-mile guard for every .docx-writing call in the app: strips
# characters XML 1.0 (and therefore .docx) never allows in text content --
# NULL/other C0 control codes, lone UTF-16 surrogates (a known artefact of
# malformed CID-keyed fonts in some tender PDFs), and the U+FFFE/U+FFFF
# non-characters. python-docx/lxml raises ValueError the instant one of
# these reaches a run, so every text write in this module and generator.py
# goes through this first regardless of whether the source (PDF extraction,
# an LLM echo, typed form input) has already been cleaned upstream.
_DOCX_ILLEGAL_CHARS_RE = re.compile("[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff￾￿]")


def docx_safe(text) -> str:
    return _DOCX_ILLEGAL_CHARS_RE.sub("", text or "")


def ensure_unique_drawing_ids(doc):
    """wp:docPr/@id must be unique across the WHOLE document package (body +
    every header/footer), not just within one part -- python-docx assigns
    these ids per-part, so a picture added to the header (the logo) and a
    picture added to the body (the wet-ink signature in the signoff block)
    can both land on docPr id="1" whenever the body hasn't had any other
    picture before that point. Word doesn't reliably render both shapes
    when their ids collide like that -- one goes invisible/unselectable --
    which is what was silently dropping the signature next to the company
    stamp on some documents (and not others, depending on how many other
    pictures happened to already be in that document's body). Call this
    once, right before doc.save(), to renumber every drawing across every
    part so no two ever collide."""
    next_id = 1
    part_roots = [doc.element.body]
    for section in doc.sections:
        part_roots.append(section.header._element)
        part_roots.append(section.footer._element)
    for root in part_roots:
        for docpr in root.findall('.//' + qn('wp:docPr')):
            docpr.set('id', str(next_id))
            next_id += 1


def set_default_font(doc, name: str, size: int):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def add_letterhead(doc):
    """Puts the logo in the section's actual Word header (not a body
    paragraph), so it repeats on every page automatically instead of only
    appearing once at the top of page 1."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Cm(6.75), height=Cm(3))

    # Header from Top fixed at 0.5cm on every page. Body content still needs
    # a gap below the header so it doesn't overlap the logo: header_distance
    # (0.5cm) + logo height (3cm) = 3.5cm before the header's content ends,
    # so top_margin is set a bit past that to leave a visible gap uniformly
    # on every page (a per-page layout property, not tied to where a
    # paragraph break happens to fall).
    section.header_distance = Cm(0.5)
    section.top_margin = Cm(4.2)


def _add_top_border(paragraph, color="1F4E79", size=6):
    """A ruled line across the top of `paragraph` -- used to visually
    separate the footer block from the document body above it."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def add_footer(doc, company: dict):
    """Company name: bold, 18pt, blue -- the whole line.
    Registered Office / Factory: only the "Registered Office:"/"Factory:"
    label itself is bold; the address value that follows it is regular
    weight. Both still 10pt black.
    Email/Phone/CIN: bold, 10pt, black -- the whole line (not split)."""
    section = doc.sections[0]
    footer = section.footer
    black = RGBColor(0x00, 0x00, 0x00)
    blue = RGBColor(0x1F, 0x4E, 0x79)

    def _new_para(first=False):
        p = footer.paragraphs[0] if first else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        return p

    def _run(p, text, bold, size, color):
        run = p.add_run(docx_safe(text))
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return run

    p = _new_para(first=True)
    _run(p, company["company_name"], True, 18, blue)

    for label, value in [
        ("Registered Office:", company["registered_office"]),
        ("Factory:", company["factory_address"]),
    ]:
        p = _new_para()
        _run(p, label + " ", True, 10, black)
        _run(p, value, False, 10, black)

    p = _new_para()
    _run(p, f"Email: {company['email']}; Phone: {company['phone']}; CIN No: {company['cin']}", True, 10, black)

    footer.paragraphs[0].paragraph_format.space_before = Pt(4)
    _add_top_border(footer.paragraphs[0])


def add_title(doc, text: str):
    # The logo now lives in the section header (repeats on every page), not
    # a body paragraph, so without a gap here the title sits flush against
    # wherever the header ends. This is the first thing added to the body,
    # so it's the one place that needs to open with a blank spacer.
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(docx_safe(text))
    run.bold = True
    doc.add_paragraph()


def add_signoff_tail(doc, company: dict, place_placeholder: str, date_placeholder: str):
    """The 'Dated / Place / For Airox... / signature-or-blank / name / title'
    block every Airox document ends with. Chained with keep_with_next /
    keep_together so Word never splits the signatory's name from their
    title across a page break."""
    p1 = doc.add_paragraph()
    p1.add_run(docx_safe(f"Dated this {date_placeholder}."))
    p2 = doc.add_paragraph()
    p2.add_run(docx_safe(f"Place: {place_placeholder}"))
    p1.paragraph_format.keep_with_next = True
    p2.paragraph_format.keep_with_next = True
    p1.paragraph_format.keep_together = True
    p2.paragraph_format.keep_together = True

    chain = []
    p_for = doc.add_paragraph()
    p_for.add_run(docx_safe(f"For {company['company_name']}")).bold = True
    chain.append(p_for)

    p_sig = doc.add_paragraph()
    if os.path.exists(STAMP_PATH) or os.path.exists(SIGNATURE_PATH):
        run = p_sig.add_run()
        if os.path.exists(SIGNATURE_PATH):
            run.add_picture(SIGNATURE_PATH, height=Inches(0.5))
        if os.path.exists(STAMP_PATH):
            p_sig.add_run("  ")
            p_sig.add_run().add_picture(STAMP_PATH, height=Inches(0.9))
    chain.append(p_sig)

    p_name = doc.add_paragraph()
    p_name.add_run(docx_safe(company["signatory_name"]))
    chain.append(p_name)

    p_title = doc.add_paragraph()
    p_title.add_run(docx_safe(company["signatory_title"])).italic = True

    for p in chain:
        p.paragraph_format.keep_with_next = True
    for p in chain + [p_title]:
        p.paragraph_format.keep_together = True
    return p_title
