"""
Step 4 of the pipeline: fill a template .docx with values from the Tender
Profile + Company Profile.

Templates use plain {{field_name}} placeholders typed directly into the
Word doc (same font/size as surrounding text). We do NOT use docxtpl/Jinja
here -- for pure key/value substitution, a direct python-docx pass is one
less dependency and keeps the merge-runs-then-replace logic visible and
auditable, which matters on documents you're signing.

Word silently splits a single visible word across multiple <w:r> runs
(spellcheck, autocorrect markers), so "{{tender_ref_no}}" you can SEE in
Word often isn't one contiguous string in the XML. We handle that by
merging a paragraph's run text, doing the replacement at paragraph level,
then writing the merged result back into the first run and clearing the
rest -- standard mail-merge technique.
"""
import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from .branding import docx_safe

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


FILL_TAG_RE = re.compile(r"(\[FILL: [a-zA-Z0-9_]+\])")

# Sentinels wrapped around any value substituted in for a {{placeholder}},
# so the segment can be tagged green below -- distinct from the yellow
# [FILL: ...] tags used when nothing was available at all. Private-Use-Area
# code points: never appear in real document text, so they're safe delimiters
# to smuggle through the plain-string regex substitution.
_FILLED_START = ""
_FILLED_END = ""
FILLED_VALUE_RE = re.compile(re.escape(_FILLED_START) + r"(.*?)" + re.escape(_FILLED_END), re.DOTALL)
SEGMENT_RE = re.compile(
    r"(\[FILL: [a-zA-Z0-9_]+\]|" + re.escape(_FILLED_START) + r".*?" + re.escape(_FILLED_END) + r")",
    re.DOTALL,
)


def _replace_in_paragraph(paragraph, values: dict) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return False

    def sub(m):
        key = m.group(1)
        val = values.get(key) or ""
        if val:
            return f"{_FILLED_START}{docx_safe(val)}{_FILLED_END}"
        return f"[FILL: {key}]"

    new_text = PLACEHOLDER_RE.sub(sub, full_text)
    if new_text == full_text:
        return False

    if not paragraph.runs:
        return False

    # Snapshot the formatting of the run we're about to overwrite, so every
    # fragment we (re)write -- filled text, [FILL: ...] tags, and plain
    # boilerplate -- keeps the surrounding sentence's font/size/bold/italic.
    # Only the substituted fragments get a highlight; the rest of the
    # sentence must not.
    base_run = paragraph.runs[0]
    base_font = base_run.font
    base = {
        "name": base_font.name,
        "size": base_font.size,
        "bold": base_run.bold,
        "italic": base_run.italic,
        "color": base_font.color.rgb if base_font.color and base_font.color.type else None,
    }

    def apply_base(run, highlight=None):
        run.font.name = base["name"]
        run.font.size = base["size"]
        run.bold = base["bold"]
        run.italic = base["italic"]
        if base["color"] is not None:
            run.font.color.rgb = base["color"]
        run.font.highlight_color = highlight

    def resolve(seg: str):
        """Strip sentinels and pick a highlight: yellow for [FILL: ...]
        (nothing found -- fill in by hand), green for a real substituted
        value (so you can tell at a glance what the system inserted vs.
        the original boilerplate wording), none for untouched template
        text."""
        if FILL_TAG_RE.fullmatch(seg):
            return seg, WD_COLOR_INDEX.YELLOW
        m = FILLED_VALUE_RE.fullmatch(seg)
        if m:
            return m.group(1), WD_COLOR_INDEX.BRIGHT_GREEN
        return seg, None

    # Clear all existing runs; we'll rebuild the paragraph's runs from the
    # segments below (this is what keeps each highlight scoped to just its
    # own fragment instead of bleeding across the whole sentence).
    for run in paragraph.runs[1:]:
        run.text = ""
    segments = [s for s in SEGMENT_RE.split(new_text) if s]
    if not segments:
        segments = [""]

    text0, highlight0 = resolve(segments[0])
    paragraph.runs[0].text = text0
    apply_base(paragraph.runs[0], highlight=highlight0)
    for seg in segments[1:]:
        text, highlight = resolve(seg)
        run = paragraph.add_run(text)
        apply_base(run, highlight=highlight)

    return True


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def fill_template(template_path: str, values: dict, output_path: str) -> list[str]:
    """values: flat dict of str -> str (already resolved, missing values
    should already be formatted as '[FILL: label]' by the caller).
    Returns list of placeholder keys that were left unresolved (not found
    in `values` at all -- distinct from ones filled with a [FILL:] tag)."""
    doc = Document(template_path)
    unresolved = set()

    all_placeholders = set()
    for p in _iter_all_paragraphs(doc):
        text = "".join(r.text for r in p.runs)
        all_placeholders.update(PLACEHOLDER_RE.findall(text))

    for p in _iter_all_paragraphs(doc):
        _replace_in_paragraph(p, values)

    unresolved = {k for k in all_placeholders if k not in values}
    doc.save(output_path)
    return sorted(unresolved)


def scan_placeholders(template_path: str) -> list[str]:
    """List every {{field}} a template references -- used to build the
    field-mapping table and to sanity-check new templates you add."""
    doc = Document(template_path)
    found = set()
    for p in _iter_all_paragraphs(doc):
        text = "".join(r.text for r in p.runs)
        found.update(PLACEHOLDER_RE.findall(text))
    return sorted(found)
