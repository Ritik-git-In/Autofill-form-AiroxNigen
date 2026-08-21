"""
Step 1 (parse) + Step 2 (canonical extract) of the pipeline.

parse_pdf_pages(path)              -> list[str], text per page, via pdfplumber
parse_reference_document(path)     -> str, flattened text of a supporting PDF/.docx
extract_profile(pages, reference_text=None)
    -> (profile: dict[str, FieldValue], required_documents: list[dict])
compose_document_text(requirement_name, tender_facts, company) -> str

extract_profile calls Claude with the FULL page-tagged text of the incoming
PDF (chunked if it's very large), plus an optional reference document's text
for facts the tender itself doesn't state, and asks for ONE structured JSON
object with two parts: the TENDER_PROFILE_FIELDS facts, and a list of every
document/certificate/annexure the tender's own checklist says the bidder
must submit (used to build the "Select documents to generate" list from the
tender itself, instead of always offering the full template library). The
prompt is explicit: copy facts verbatim, and leave a field null if it isn't
confidently present in the text -- never infer, estimate, or reword. Each
required document also comes back with has_own_content: true if the tender
itself bundles that document's actual text somewhere (so it can be
reproduced verbatim -- see generator.build_verbatim_document), or false if
it's known only from a checklist mention with nothing to copy from
anywhere in the tender (see compose_document_text below, which drafts that
one from scratch instead).

Uses whichever LLM backend has an API key configured -- KIMI_API_KEY
(Moonshot AI's Kimi models, via their OpenAI-compatible endpoint) is tried
first, then ANTHROPIC_API_KEY (Claude). With neither set, falls back to a
small regex-based extractor so the rest of the app (review UI, filling,
generation) can be built and tested without an API key. It can't sensibly
detect a required-documents list, so it always returns an empty one -- the
caller falls back to offering the full template library in that case.
compose_document_text has no such fallback -- drafting original content
needs a real LLM, so it raises if neither API key is set (unreachable in
practice, since required_documents is always empty without one anyway).
"""
import json
import os
import re
import threading
from concurrent.futures import ThreadPoolExecutor

from .schema import FieldValue, TENDER_PROFILE_FIELDS, new_tender_profile

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# pdfplumber's extract_text() is pure Python, but enough of its actual work
# (zlib decompression, font/encoding parsing) releases the GIL that running
# pages through a thread pool is a real win, not just overhead -- measured
# ~3.5x faster on a 191-page tender (21s -> 6s). Capped at 8 so this doesn't
# oversubscribe on smaller machines.
_PARSE_WORKERS = min(os.cpu_count() or 4, 8)

# Characters XML 1.0 (and therefore .docx, which is XML under the hood)
# does not allow anywhere in text content: NULL/other C0 control codes
# other than tab/newline/carriage-return, lone UTF-16 surrogates, and the
# U+FFFE/U+FFFF non-characters. Some tenders' PDFs have these embedded --
# a malformed/CID-keyed font whose glyph mapping pdfminer can't resolve
# cleanly is a common source of lone surrogate codepoints in particular --
# and python-docx/lxml raises ValueError the moment such text is written
# into a run, so this is stripped once at the source rather than patched
# at every place that later writes PDF-derived text into a .docx.
_XML_ILLEGAL_CHARS_RE = re.compile(
    "[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff￾￿]"
)


def _clean_text(text: str) -> str:
    return _XML_ILLEGAL_CHARS_RE.sub("", text)


def _safe_extract_text(page) -> str:
    """Some real-world tender PDFs have a page with a malformed embedded
    object -- e.g. a colour space definition pdfminer can't parse -- that
    isn't the rest of the document's fault. Treat that one page as empty
    (same as a scanned/image-only page) instead of taking down extraction
    for the whole 190-page tender over one bad page."""
    try:
        return _clean_text(page.extract_text() or "")
    except Exception:
        return ""


def _run_per_page(path: str, page_indices, worker):
    """Run worker(page) across the given 0-indexed page positions in a
    thread pool, giving each worker THREAD its own independent
    pdfplumber.PDF for `path` instead of sharing one pdfplumber.open()
    across threads.

    Sharing one PDF object's .pages between threads looked safe (no
    exception, ever) but wasn't: pdfminer's underlying PDFDocument /
    PDFResourceManager keep parsing caches (fonts, xobjects, content-stream
    state) that aren't safe under concurrent access from different pages.
    Measured directly against a real 187-page tender: about 1 run in 3,
    a whole contiguous run of pages came back with empty text AND zero
    tables -- no exception raised, so _safe_extract_text's try/except
    never even saw anything wrong, it just silently produced an empty
    page. That's much worse than a crash: the generated .docx still shows
    a "[Tender page N]" caption for every page (so it looks complete at a
    glance) while whole checklist/annex pages are blank underneath.
    Giving every worker thread its own PDF parse removes the shared state
    entirely, so a thread can never observe another thread's mid-parse
    cache. Opened once per thread (not per page) so this doesn't reintroduce
    the per-page-reopen overhead that would undo the earlier threading
    speed-up.
    """
    local = threading.local()
    opened = []
    opened_lock = threading.Lock()

    def _get_pdf():
        pdf = getattr(local, "pdf", None)
        if pdf is None:
            pdf = pdfplumber.open(path)
            local.pdf = pdf
            with opened_lock:
                opened.append(pdf)
        return pdf

    def _run(i):
        return worker(_get_pdf().pages[i])

    try:
        with ThreadPoolExecutor(max_workers=_PARSE_WORKERS) as ex:
            return list(ex.map(_run, page_indices))
    finally:
        for pdf in opened:
            pdf.close()


def parse_pdf_pages(path: str) -> list[str]:
    """Return a list of page texts. Pages that come back empty (likely
    scanned/image-only, or an unparseable embedded object) are flagged so
    the UI can warn OCR may be needed."""
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is not installed. pip install pdfplumber")
    with pdfplumber.open(path) as pdf:
        n = len(pdf.pages)
    return _run_per_page(path, range(n), _safe_extract_text)


def parse_pdf_page_range_detailed(path: str, start: int, end: int) -> list[dict]:
    """Like parse_pdf_page_range, but also returns any genuine tables on each
    page (pdfplumber's grid-line based extract_tables()) alongside the plain
    text. Verbatim-reproducing a tender page as flat paragraph text badly
    scrambles anything with side-by-side columns -- MRPL's Inspection & Test
    Plan checklists and QAP forms come out as unreadable word-soup that way,
    because plain-text extraction has no notion of column boundaries. Where
    pdfplumber finds a real table on the page, the caller renders it as an
    actual Word table (columns preserved) instead of flattening it."""
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is not installed. pip install pdfplumber")

    def _extract_one(page):
        try:
            tables = page.extract_tables() or []
        except Exception:
            tables = []
        cleaned_tables = [
            [[_clean_text(cell or "") for cell in row] for row in table]
            for table in tables if table
        ]
        return {"page": page.page_number, "text": _safe_extract_text(page), "tables": cleaned_tables}

    with pdfplumber.open(path) as pdf:
        lo = max(1, start)
        hi = min(len(pdf.pages), end)
    return _run_per_page(path, range(lo - 1, hi), _extract_one)


def parse_docx_text(path: str) -> str:
    """Flatten a .docx's paragraphs and table cells into plain text."""
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def parse_reference_document(path: str) -> str:
    """Read a supporting reference document (one of Airox's own previously
    filled documents) into plain text, regardless of whether it's a PDF or
    a .docx. Used only as extra context for extraction -- never filled
    into a template directly."""
    lower = path.lower()
    if lower.endswith(".pdf"):
        return "\n\n".join(parse_pdf_pages(path))
    if lower.endswith(".docx"):
        return parse_docx_text(path)
    raise ValueError(f"Unsupported reference document type: {path}")


def _build_prompt(pages: list[str], reference_text: str | None = None) -> str:
    tagged = []
    for i, text in enumerate(pages, start=1):
        if text.strip():
            tagged.append(f"--- PAGE {i} ---\n{text}")
    body = "\n\n".join(tagged)

    field_lines = "\n".join(f'    "{key}": "{desc}"' for key, desc in TENDER_PROFILE_FIELDS)

    reference_block = ""
    if reference_text and reference_text.strip():
        reference_block = f"""

You may also use the REFERENCE MATERIAL below -- one or more of Airox's own
previously prepared documents (each marked with its own "--- REFERENCE
FILE: ... ---" header) -- for facts the tender itself doesn't state (e.g.
Airox's own net worth or business address). Always prefer the tender
document when both mention the same fact. A fact sourced only from the
reference material has no page number in the tender -- leave source_page
null for those.

REFERENCE MATERIAL:
{reference_text}
"""

    return f"""You are reading a tender/RFP document for Airox Nigen Equipments Pvt.
Ltd. to prepare their bid submission.

Return ONLY a JSON object with exactly two top-level keys: "fields" and
"required_documents".

"fields" -- an object with exactly these keys. For each key, return an
object {{"value": <string or null>, "source_page": <int or null>, "confidence": "high"|"low"}}.
{{
{field_lines}
}}

"required_documents" -- a JSON array listing every distinct document,
certificate, affidavit, annexure, form, specification, plan, or procedure
this tender requires the bidder to submit, sign, or countersign as part of
the bid. Two sources of these:
  (a) an explicit "List of Documents", "Checklist of Documents",
      "Annexures", "Appendices", or similar section that enumerates what
      must be submitted; and
  (b) standalone documents bundled inside the tender package itself that
      the bidder must sign/countersign -- recognisable because each one
      starts on its own title page and every page of it carries a line
      like "Tenderer's/Bidder's Signature with seal" or an equivalent
      acknowledgement line. Treat each such bundled document (e.g. an
      Inspection & Test Plan, Quality Assurance Plan, List of Approved
      Makes, technical specification chapter) as one entry, using its own
      title as "name" -- do NOT include the tender's overall front-matter,
      or documents that are the Authority's own already-signed/issued
      material (e.g. a safety policy circulated for compliance, not for
      the bidder to sign), or pricing/rate schedules submitted separately
      via the e-portal.
For each one, return an object:
  {{"name": <exact name/title as stated in the tender>, "source_page": <int or null>, "end_page": <int or null>, "has_own_content": true|false}}
"has_own_content" is true ONLY if this document exists as its own bundled
pages somewhere in the tender per rule (b) above (its own title page, a
per-page signature/seal line) -- meaning its exact wording can be copied
verbatim. It is false if you know about this document only because it's
named in a checklist per rule (a), with no dedicated pages of its own
anywhere else in the tender -- there is nothing to copy for it, it would
have to be drafted from scratch. A single requirement can appear in a
checklist AND be bundled elsewhere (has_own_content: true, pointing at the
bundled pages, not the checklist line) -- check the whole document before
deciding false.
"source_page" is the page where this document's title/heading first
appears (only meaningful when has_own_content is true -- when false it may
point at the checklist page mentioning it, or be null). "end_page" is the
last page belonging to this same document (the page immediately before the
next document/section starts, or the tender's last page if it runs to the
end) -- give your best confident estimate when has_own_content is true;
leave it null only if you genuinely cannot tell, and always null when
has_own_content is false.
Use the tender's own wording for "name" -- do not paraphrase or invent one.
If the same document is referenced more than once, list it only once. If you
cannot find any such documents, return an empty array rather than guessing.

Rules:
- Copy the exact wording found in the source text -- do not paraphrase,
  reword, correct, or re-punctuate a fact or a document name. If you cannot
  copy it verbatim with confidence, treat it as not found.
- If a fact is not clearly and confidently present in the text, set its
  value to null. NEVER estimate, guess, or infer a value that isn't
  actually stated.
- source_page is the PAGE number (from the "--- PAGE N ---" markers) where
  you found it.
- confidence "low" if you're inferring from indirect context, "high" if
  it's stated directly.

TENDER DOCUMENT TEXT:
{body}
{reference_block}"""


def _as_str(value):
    """The prompt asks the model for plain string values throughout, but it
    occasionally nests one anyway (e.g. a dict for a compound fact instead
    of a flat string). That's not something the rest of the app can use as
    a Tender Profile fact or document name, and previously crashed
    _clean_text outright (it only knows how to regex-scrub a str). Treat a
    wrong-shaped value the same as the model returning null -- "not
    confidently present" -- rather than crashing or guessing a
    stringification of whatever shape it sent back."""
    return value if isinstance(value, str) else None


def _parse_llm_response(text: str) -> tuple[dict, list[dict]]:
    """Shared response parser for any backend: strips a markdown code fence
    if the model wrapped its JSON in one, then builds the profile +
    required_documents from the {"fields": ..., "required_documents": ...}
    shape every backend is prompted to return."""
    cleaned = re.sub(r"^```json|```$", "", text.strip(), flags=re.MULTILINE).strip()
    data = json.loads(cleaned)

    fields_data = data.get("fields") or {}
    profile = new_tender_profile()
    for key, _ in TENDER_PROFILE_FIELDS:
        raw = fields_data.get(key) or {}
        value = _as_str(raw.get("value"))
        profile[key] = FieldValue(
            value=_clean_text(value) if value else value,
            source_page=raw.get("source_page"),
            confidence=raw.get("confidence", "unset") if value else "unset",
        )

    required_documents = []
    for d in (data.get("required_documents") or []):
        name = (_as_str(d.get("name")) or "").strip()
        if not name:
            continue
        required_documents.append({
            "name": _clean_text(name),
            "source_page": d.get("source_page"),
            "end_page": d.get("end_page"),
            "has_own_content": d.get("has_own_content"),
        })

    return profile, required_documents


def extract_profile_llm(
    pages: list[str], reference_text: str | None = None, model: str = "claude-sonnet-4-6"
) -> tuple[dict, list[dict]]:
    """Real extraction path using the Anthropic API. Requires ANTHROPIC_API_KEY."""
    import anthropic

    client = anthropic.Anthropic()
    prompt = _build_prompt(pages, reference_text)

    # crude guard: if the doc is huge, this single call may exceed context.
    # For very large RFPs, chunk pages and merge, keeping the first non-null
    # + highest-confidence value per field. Left as a follow-up; 200 pages
    # of typical RFP text is usually within a single call's budget.
    resp = client.messages.create(
        model=model,
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )
    text = "".join(b.text for b in resp.content if getattr(b, "type", None) == "text")
    return _parse_llm_response(text)


def extract_profile_kimi(pages: list[str], reference_text: str | None = None, model: str | None = None) -> tuple[dict, list[dict]]:
    """Real extraction path using Moonshot AI's Kimi models, over their
    OpenAI-compatible /chat/completions endpoint. Requires KIMI_API_KEY.

    Base URL and model are configurable via env vars because Moonshot runs
    both a China endpoint (api.moonshot.cn) and an international one
    (api.moonshot.ai), and exact model ids change over time -- if the
    default model below comes back with a "model not found" error, check
    your Kimi console for the right id and set KIMI_MODEL in .env (same for
    KIMI_BASE_URL if your account is on the .cn platform)."""
    import requests

    base_url = os.environ.get("KIMI_BASE_URL", "https://api.moonshot.ai/v1").rstrip("/")
    model = model or os.environ.get("KIMI_MODEL", "moonshot-v1-128k")
    prompt = _build_prompt(pages, reference_text)

    resp = requests.post(
        f"{base_url}/chat/completions",
        headers={
            "Authorization": f"Bearer {os.environ['KIMI_API_KEY']}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2,
            "max_tokens": 4000,
        },
        timeout=120,
    )
    resp.raise_for_status()
    text = resp.json()["choices"][0]["message"]["content"]
    return _parse_llm_response(text)


# ---------------------------------------------------------------------------
# AI-drafting: for a required document that has no hand-authored template
# AND no bundled text anywhere in the tender to reproduce (has_own_content
# is false in required_documents -- see the prompt above). Everything the
# tender itself provides text for is still reproduced verbatim, never
# rewritten; this path only exists for the genuine gap where nothing exists
# to copy from at all.
# ---------------------------------------------------------------------------
def _compose_prompt(requirement_name: str, tender_facts: dict, company: dict) -> str:
    facts_lines = "\n".join(f"- {k}: {v}" for k, v in tender_facts.items() if v)
    return f"""You are drafting a document titled "{requirement_name}" that
{company['company_name']} must submit as part of a tender bid. The tender
requires this document by name (in its checklist of documents to submit)
but does not provide a ready-made format for it anywhere in the tender
package -- you must compose it from scratch, in the standard professional
style used for this kind of declaration, certificate, undertaking, or
annexure in Indian government/PSU tenders.

Known facts you may use (do not introduce any fact not given here):
{facts_lines or "(no specific tender facts were extracted -- keep the content generic to this document type, and rely on [FILL: ...] placeholders below for anything specific)"}

Company details:
- Company name: {company['company_name']}
- Registered office: {company['registered_office']}
- CIN: {company['cin']}
- GSTIN: {company['gstin']}

Rules:
- Write ONLY the body of the document -- do not repeat the title, and do
  not add a letterhead, a "Dated"/"Place" line, or a signature block
  ("For {company['company_name']}", signatory name/title); all of that is
  added separately, after your text.
- Match the tone and structure a real Indian tender document of this type
  would have: formal declaration/undertaking language, numbered clauses
  where that's the convention for this document type.
- If a specific fact is needed (a reference/registration number, an amount,
  a date, a clause number from the tender) that isn't given above, write
  "[FILL: <what's needed>]" in its place instead of inventing a
  plausible-looking value. Never fabricate a specific number, date, name,
  or fact -- an invented one is worse than a visible placeholder.
- Return plain text only: paragraphs separated by a blank line, no
  markdown formatting (no `#`, `**`, bullet characters), no JSON.
"""


def compose_document_text(requirement_name: str, tender_facts: dict, company: dict) -> str:
    """AI-drafts the body text for a required tender document that has no
    template and no bundled tender text to reproduce (the has_own_content:
    false case from extract_profile's required_documents). Requires
    KIMI_API_KEY or ANTHROPIC_API_KEY -- there's no offline fallback for
    composing original content the way there is for extraction."""
    prompt = _compose_prompt(requirement_name, tender_facts, company)
    if os.environ.get("KIMI_API_KEY"):
        return _call_kimi_text(prompt)
    if os.environ.get("ANTHROPIC_API_KEY"):
        return _call_anthropic_text(prompt)
    raise RuntimeError(
        "No LLM API key configured (KIMI_API_KEY or ANTHROPIC_API_KEY) -- "
        "AI-drafting a missing document requires one."
    )


def _call_kimi_text(prompt: str, model: str | None = None) -> str:
    import requests

    base_url = os.environ.get("KIMI_BASE_URL", "https://api.moonshot.ai/v1").rstrip("/")
    model = model or os.environ.get("KIMI_MODEL", "moonshot-v1-128k")
    resp = requests.post(
        f"{base_url}/chat/completions",
        headers={
            "Authorization": f"Bearer {os.environ['KIMI_API_KEY']}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 2000,
        },
        timeout=120,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def _call_anthropic_text(prompt: str, model: str = "claude-sonnet-4-6") -> str:
    import anthropic

    client = anthropic.Anthropic()
    resp = client.messages.create(
        model=model,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(b.text for b in resp.content if getattr(b, "type", None) == "text").strip()


# ---------------------------------------------------------------------------
# Offline fallback: simple regex extractor, used when there is no API key.
# This is intentionally narrow -- it exists so the pipeline is testable end
# to end without network/API access, NOT as a production extraction method.
# It never attempts required-document detection (that needs real reading
# comprehension) -- it always returns an empty list for that.
# ---------------------------------------------------------------------------
def extract_profile_regex(pages: list[str]) -> tuple[dict, list[dict]]:
    full = "\n".join(f"[[PAGE {i+1}]]\n{t}" for i, t in enumerate(pages))
    profile = new_tender_profile()

    def find(pattern, group=1, flags=re.IGNORECASE):
        m = re.search(pattern, full, flags)
        if not m:
            return None, None
        page_m = re.search(r"\[\[PAGE (\d+)\]\][^\[]*" + re.escape(m.group(0)), full)
        page = int(page_m.group(1)) if page_m else None
        return m.group(group).strip(), page

    patterns = {
        "tender_ref_no": r"Tender\s*/\s*Bid Enquiry No\.?\s*:?\s*([A-Z0-9/\-\.]+)",
        "tender_ref_date": r"Bid Enquiry No\.?[^\n]*dated\s*([\d\-]+)",
        "emd_amount": r"(Rs\.?\s*[\d,]+)\s*\)?\s*(?:\[Bid Security\]|Bid Security|EMD)",
    }
    for key, pat in patterns.items():
        val, page = find(pat)
        if val:
            profile[key] = FieldValue(value=val, source_page=page, confidence="low")

    # Project title: look for the long quoted "Selection of EPC..." caption
    m = re.search(r'["“]Selection of[^"”]{20,300}', full, re.IGNORECASE)
    if m:
        page_m = re.search(r"\[\[PAGE (\d+)\]\][^\[]*" + re.escape(m.group(0)[:40]), full)
        cleaned = re.sub(r"\s+", " ", m.group(0)).strip(' "“”')
        profile["project_title"] = FieldValue(
            value=cleaned,
            source_page=int(page_m.group(1)) if page_m else None,
            confidence="low",
        )

    m = re.search(r"(\d+\s*kW)", full)
    if m:
        profile["project_capacity"] = FieldValue(value=m.group(1), confidence="low")

    m = re.search(r"within\s+([A-Z][A-Za-z0-9 ]+Premises)", full)
    if m:
        profile["project_site"] = FieldValue(value=m.group(1).strip(), confidence="low")

    m = re.search(r"for\s+(\d+)\s+years", full, re.IGNORECASE)
    if m:
        profile["contract_period"] = FieldValue(value=f"{m.group(1)} years", confidence="low")

    return profile, []


def extract_profile(pages: list[str], reference_text: str | None = None) -> tuple[dict, list[dict]]:
    if os.environ.get("KIMI_API_KEY"):
        return extract_profile_kimi(pages, reference_text=reference_text)
    if os.environ.get("ANTHROPIC_API_KEY"):
        return extract_profile_llm(pages, reference_text=reference_text)
    return extract_profile_regex(pages)
