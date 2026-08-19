"""
Builds the initial templates_store/*.docx library from Airox's real,
previously-submitted documents (content supplied by the user), with the
changing facts swapped for {{placeholders}}.

Run once to seed the template library:
    python build_templates.py

You will replace/add to these over time as you bring in more of your real
templates -- this script is just how the starter set was produced, it's
not something the running app depends on.
"""
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
STORE = os.path.join(HERE, "templates_store")
LOGO = os.path.join(HERE, "static", "airox_logo.jpg")
COMPANY = json.load(open(os.path.join(HERE, "config", "company_profile.json")))

os.makedirs(STORE, exist_ok=True)


def set_default_font(doc, name="Times New Roman", size=11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def add_letterhead(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(LOGO):
        run.add_picture(LOGO, width=Inches(1.6))
    section.top_margin = Inches(1.5)  # keep gap between header logo and body on every page


def _add_top_border(paragraph, color="1F4E79", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def add_footer(doc):
    # Company name: bold, 18pt, blue -- whole line. Registered Office /
    # Factory: only the label is bold, the address value after it is
    # regular weight (both still 10pt black). Email/Phone/CIN: bold, 10pt,
    # whole line.
    section = doc.sections[0]
    footer = section.footer
    black = RGBColor(0x00, 0x00, 0x00)
    blue = RGBColor(0x1F, 0x4E, 0x79)

    def _new_para(first=False):
        p = footer.paragraphs[0] if first else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Kill Word's default "space before/after" on paragraphs -- this is
        # what was showing up as an oversized gap between the Registered
        # Office line and the Factory line.
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        return p

    def _run(p, text, bold, size, color):
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color

    p = _new_para(first=True)
    _run(p, COMPANY["company_name"], True, 18, blue)

    for label, value in [
        ("Registered Office:", COMPANY["registered_office"]),
        ("Factory:", COMPANY["factory_address"]),
    ]:
        p = _new_para()
        _run(p, label + " ", True, 10, black)
        _run(p, value, False, 10, black)

    p = _new_para()
    _run(p, f"Email: {COMPANY['email']}; Phone: {COMPANY['phone']}; CIN No: {COMPANY['cin']}", True, 10, black)

    footer.paragraphs[0].paragraph_format.space_before = Pt(4)
    _add_top_border(footer.paragraphs[0])


def add_title(doc, text):
    doc.add_paragraph()  # gap between the header logo and the title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    doc.add_paragraph()


def add_para(doc, text, bold=False, space_after=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if not space_after:
        p.paragraph_format.space_after = Pt(0)
    return p


def add_signoff_tail(doc, include_for_company=True):
    """The 'For Airox Nigen... / [blank] / signatory name / title' chunk
    that ends every document, identically. Chained with keep_with_next /
    keep_together so Word can never split the signatory's name from their
    title across a page break."""
    chain = []
    if include_for_company:
        p = add_para(doc, f"For {COMPANY['company_name']}", bold=True)
        chain.append(p)
    p2 = doc.add_paragraph()
    p3 = add_para(doc, COMPANY["signatory_name"])
    chain.extend([p2, p3])
    p4 = doc.add_paragraph()
    run = p4.add_run(COMPANY["signatory_title"])
    run.italic = True
    for p in chain:
        p.paragraph_format.keep_with_next = True
    for p in chain + [p4]:
        p.paragraph_format.keep_together = True
    return p4


def add_signature_block(doc, place_placeholder="{{signing_place}}", date_placeholder="{{signing_date}}"):
    """The standard sign-off block used at the end of most Airox documents:
    Dated/Place lines + the signoff tail. All chained with keep_with_next so
    Word can never break the page between, say, the signatory's name and
    their title -- that's what was splitting 'Anil Kumar Agrawal' from
    'Managing Director' across two pages."""
    p1 = add_para(doc, f"Dated this {date_placeholder}.")
    p2 = add_para(doc, f"Place: {place_placeholder}")
    p1.paragraph_format.keep_with_next = True
    p2.paragraph_format.keep_with_next = True
    p1.paragraph_format.keep_together = True
    p2.paragraph_format.keep_together = True
    add_signoff_tail(doc)


def keep_with_next(*paragraphs):
    for p in paragraphs:
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True


def ensure_unique_drawing_ids(doc):
    """wp:docPr/@id must be unique across the whole package (body + every
    header/footer), not just within one part -- python-docx assigns these
    per-part, so the header logo and a body picture (signature/stamp) can
    both land on id="1" and Word then only renders one of them. Call right
    before doc.save()."""
    next_id = 1
    part_roots = [doc.element.body]
    for section in doc.sections:
        part_roots.append(section.header._element)
        part_roots.append(section.footer._element)
    for root in part_roots:
        for docpr in root.findall('.//' + qn('wp:docPr')):
            docpr.set('id', str(next_id))
            next_id += 1


def base_doc():
    doc = Document()
    set_default_font(doc)
    add_letterhead(doc)
    add_footer(doc)
    return doc


# ---------------------------------------------------------------------
# 1. Anti-Collusion Certificate
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "ANTI-COLLUSION CERTIFICATE")
add_para(doc,
    'We hereby certify and confirm that in the preparation and submission of our Bid for '
    '"{{project_title}}" District of {{project_state}}, we have not acted in concert or in '
    'collusion with any other Bidder or other person(s) and also not done any act, deed or thing '
    'which is or could be regarded as anticompetitive.')
add_para(doc,
    'We further confirm that we have not offered nor will offer any illegal gratification in cash '
    'or kind to any person or agency in connection with the instant Bid.')
doc.add_paragraph()
add_signature_block(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "anti_collusion_certificate.docx"))

# ---------------------------------------------------------------------
# 2. Anti-Blacklisting Affidavit
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "ANTI-BLACKLISTING AFFIDAVIT")
add_para(doc,
    f'We, M/s. {COMPANY["company_name"]} having its registered office {COMPANY["registered_office"]}, '
    'hereby certify and confirm that we or any of our promoter/s / director/s are not barred by '
    '{{authority_name}} / any other entity of Government of {{project_state}} or blacklisted by any '
    'state government or central government / department / agency in India from participating in '
    'Project/s, either individually or as member of a Consortium as on {{signing_date}}.')
add_para(doc,
    'We further confirm that we are aware that our Bid for the Project would be liable for rejection '
    'in case any material misrepresentation is made or discovered with regard to the requirements of '
    'this RFP at any stage of the Bidding Process or thereafter during the agreement period.')
doc.add_paragraph()
add_signature_block(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "anti_blacklisting_affidavit.docx"))

# ---------------------------------------------------------------------
# 3. Statement of Legal Capacity
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "STATEMENT OF LEGAL CAPACITY")
add_para(doc, "Ref. Date: {{signing_date}}", space_after=True)
doc.add_paragraph()
add_para(doc, "To,")
add_para(doc, "Managing Director")
add_para(doc, "{{authority_name}}")
doc.add_paragraph()
add_para(doc, "Dear Sir/ Madam,")
add_para(doc,
    f'We hereby confirm that we {COMPANY["company_name"]} satisfy the terms and conditions laid out '
    'in the RFP document.')
add_para(doc,
    f'We have agreed that {COMPANY["signatory_name"]} will act as our representative and has been '
    'duly authorized to submit to the RFP. Further, the authorized signatory is vested with requisite '
    'powers to furnish such letter and authenticate the same.')
doc.add_paragraph()
add_para(doc, "Thanking you,")
add_para(doc, "Yours faithfully,")
add_signoff_tail(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "statement_of_legal_capacity.docx"))

# ---------------------------------------------------------------------
# 4. Restrictions on Sourcing of Equipment (land-border vendor declaration)
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc,
    "RESTRICTIONS ON SOURCING OF EQUIPMENTS/MATERIALS BY THE BIDDER FROM A VENDOR OF A "
    "COUNTRY WHICH SHARES A LAND BORDER WITH INDIA")
add_para(doc, "Tender / Bid Enquiry No.: {{tender_ref_no}} dated {{tender_ref_date}}.", bold=True)
add_para(doc,
    f'We hereby declare that {COMPANY["company_name"]} is an Indian company and is not a bidder '
    'from a country sharing a land border with India as per the applicable Government of India '
    'guidelines.')
add_para(doc,
    'I have read the clause regarding restrictions on procurement from a bidder/vendor of a country '
    'which shares a land border with India. I certify that we are not from such a country or have '
    'been registered with the Competent Authority. I hereby certify that we fulfill all requirements '
    'in this regard and is eligible to be considered.')
doc.add_paragraph()
add_para(doc, "Yours faithfully,")
doc.add_paragraph()
add_signature_block(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "restrictions_on_sourcing.docx"))

# ---------------------------------------------------------------------
# 5. Certification for not availing subsidy/grant
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc,
    "CERTIFICATION FOR NOT AVAILING/ TO AVAIL SUBSIDY/ GRANT/ CENTRAL FINANCIAL ASSISTANCE "
    "FROM MNRE AND/OR ANY OTHER STATE GOVERNMENT AND/OR ANY OTHER CENTRAL GOVERNMENT")
add_para(doc,
    f'This is to certify that we, {COMPANY["company_name"]} have applied for the installation of '
    '{{project_title}} to be implemented at a location within the geographical boundary of the '
    'respective Talukas/ Legislative Constituencies of {{project_state}} for participating under '
    'this RFP and we have not availed any subsidy/ grant/ central financial assistance from Ministry '
    'of New and Renewable Energy (MNRE) and/or any other State Government and/or any other Central '
    'Government, in this regard.')
add_para(doc,
    'We hereby also agree not to claim any subsidy/grant/ central financial assistance either from '
    'MNRE and/or any other State Government and/or any other Central Government for the same. In '
    'case it is found that we have availed any subsidy/ grant/ central financial assistance from '
    'MNRE and/or any other State Government and/or any other Central Government then {{authority_short}}, '
    'at its own discretion, can cancel all the capacities quoted in our Bids and the Bid Security/ '
    'Performance Security (if any) as submitted by us can also be forfeited. Further, {{authority_short}} '
    'can blacklist us for participating in any further Bid in {{authority_short}} up to the level of '
    'participating in any tender for Government of {{project_state}}.')
doc.add_paragraph()
add_signature_block(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "certification_not_availing_subsidy.docx"))

# ---------------------------------------------------------------------
# 6. Declaration of Shareholding Pattern (company master data table, not
#    extracted from the incoming PDF -- shown here for completeness)
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "DECLARATION OF SHAREHOLDING PATTERN OF THE BIDDER")
add_para(doc,
    'We hereby declare information of all the entities holding 15% or more shareholding in the '
    'Bidder, directly/indirectly. The information includes any compulsorily convertible Preference '
    'Shares and/or Debentures, a declaration of the likely shareholding after conversion of such '
    'instruments. The information provided herein clearly indicates the foreign shareholding and '
    'domestic shareholding in the Bidder (Differentiate between Foreign Shareholding and Domestic '
    'Shareholding):')
add_para(doc, "(Information on shareholding not more than seven (7) days from the date of submission of Technical Bid).")
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, h in enumerate(["SI. No", "Name of the Shareholder", "Domestic Share holder", "Foreign Share holder"]):
    hdr[i].text = h
for idx, sh in enumerate(COMPANY["shareholders"], start=1):
    row = table.add_row().cells
    row[0].text = str(idx)
    row[1].text = sh["name"]
    row[2].text = sh["domestic"]
    row[3].text = sh["foreign"]
doc.add_paragraph()
add_para(doc, "Yours faithfully,")
add_signature_block(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "declaration_of_shareholding.docx"))

# ---------------------------------------------------------------------
# 7. Form of Tender
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "FORM OF TENDER")
add_para(doc, "To,")
add_para(doc, "Managing Director")
add_para(doc, "{{authority_name}}")
add_para(doc, "{{authority_address}}")
doc.add_paragraph()
add_para(doc, "Sir/Madam,", bold=True)
add_para(doc,
    'We offer to execute the works described above in accordance with the Conditions of Contract '
    'accompanying this Tender with the Contract Price quoted in item-wise Financial Offer in '
    'e-procurement portal.')
add_para(doc,
    'This Tender and your written acceptance of it shall constitute a binding contract between us. '
    'We understand that you are not bound to accept the lowest or any tender you receive.')
add_para(doc,
    'We undertake that, in competing for (and, if the award is made to us, in executing) the above '
    'contract, we will strictly observe the laws against fraud and corruption in force in India '
    'namely "Prevention of Corruption Act 1988".')
add_para(doc,
    'We hereby confirm that this Tender complies with the Tender validity and Earnest Money Deposit '
    'required by the Tender documents.')
add_para(doc, 'We attach herewith our current income-tax certificate ({{financial_year}}).')
doc.add_paragraph()
p_yf = add_para(doc, "Yours faithfully,")
tail_title = add_signoff_tail(doc, include_for_company=False)
p_tenderer = add_para(doc, f"Name of Tenderer: {COMPANY['company_name']}")
p_address = add_para(doc, f"Address: {COMPANY['registered_office']}")
p_gstin = add_para(doc, f"GSTIN: {COMPANY['gstin']}")
keep_with_next(p_yf, tail_title, p_tenderer, p_address)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "form_of_tender.docx"))

# ---------------------------------------------------------------------
# 8. Power of Attorney for Signing of Bid
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "POWER OF ATTORNEY FOR SIGNING OF BID")
add_para(doc,
    f'Know all men by these presents, we M/s. {COMPANY["company_name"]} having its registered office '
    f'{COMPANY["registered_office"]}, do hereby irrevocably constitute, nominate, appoint and authorise '
    f'{COMPANY["signatory_name"]}, son of {COMPANY["signatory_father_name"]} and presently residing at '
    f'{COMPANY["signatory_residence"]} who is presently employed with us and holding the position of '
    f'{COMPANY["signatory_title"]}, as our true and lawful attorney (hereinafter referred to as the '
    '"Attorney") to do in our name and on our behalf, all such acts, deeds and things as are necessary '
    'or required in connection with or incidental to submission of our Bid for the "{{project_title}}" '
    'as per RFP no. {{tender_ref_no}} dated {{tender_ref_date}}, proposed by the {{authority_name}}, '
    'including but not limited to signing and submission of all applications, Bids and other documents '
    'and writings, participate in Bidders\' and other conferences and providing information / '
    'responses to {{authority_short}}, representing us in all matters before {{authority_short}}, '
    'signing and execution of all contracts including undertakings consequent to acceptance of our '
    'Bid, and generally dealing with {{authority_short}} in all matters in connection with or relating '
    'to or arising out of our Bid for the said Project and/or upon award thereof to us and/or till '
    'completion of contract.')
add_para(doc,
    'AND we hereby agree to ratify and confirm and do hereby ratify and confirm all acts, deeds and '
    'things lawfully done or caused to be done by our said Attorney pursuant to and in exercise of the '
    'powers conferred by this Power of Attorney and that all acts, deeds and things done by our said '
    'Attorney in exercise of the powers hereby conferred shall and shall always be deemed to have been '
    'done by us.')
add_para(doc, f'IN WITNESS WHEREOF WE {COMPANY["company_name"].upper()}, THE ABOVE NAMED PRINCIPAL '
    'HAVE EXECUTED THIS POWER OF ATTORNEY ON THIS {{signing_date}}.')
doc.add_paragraph()
p_yf = add_para(doc, "Yours faithfully,")
p_place = add_para(doc, "Place: {{signing_place}}")
keep_with_next(p_yf, p_place)
add_signoff_tail(doc)
doc.add_paragraph()
add_para(doc, "Witnesses:")
add_para(doc, "1.")
add_para(doc, "2.")
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "power_of_attorney.docx"))

# ---------------------------------------------------------------------
# 9. Site Details -- the one template with genuinely site-specific facts
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "SITE DETAILS")
add_para(doc,
    'We hereby certify that we agree to proceed with the project at the location mentioned below and '
    'that the information provided is true and correct to the best of our knowledge and belief.')
table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
rows = [
    ("Site Location", "{{project_site}}"),
    ("Taluk", "{{site_taluk}}"),
    ("District", "{{project_district}}"),
    ("State", "{{project_state}}"),
    ("Co-ordinates", "{{site_coordinates}}"),
    ("Site Coordinator Contact No.", "{{site_contact_name}} Ph: {{site_contact_phone}}"),
]
for label, val in rows:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = val
doc.add_paragraph()
p_yf = add_para(doc, "Yours faithfully,")
p_place = add_para(doc, "Place: {{signing_place}}")
p_date = add_para(doc, "Date: {{signing_date}}")
keep_with_next(p_yf, p_place, p_date)
add_signoff_tail(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "site_details.docx"))

# ---------------------------------------------------------------------
# 10. Financial Capacity -- mostly Airox's OWN company financials, not
#     extracted from the incoming PDF. Net worth cell is left blank on
#     purpose (a manual figure your accountant confirms per bid).
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "FINANCIAL CAPACITY")
table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Bidder Type"
table.rows[0].cells[1].text = "Net Worth (INR Crores)"
table.rows[1].cells[0].text = "Single Business Entity Bidder"
table.rows[1].cells[1].text = "{{net_worth_inr_crores}}"
doc.add_paragraph()
add_para(doc,
    'The Bidder shall attach copies of the balance sheets, financial statements and Annual Reports as '
    'on March 31, {{financial_statements_year}}, audited by a Chartered Accountant, complete with all '
    'notes, corresponding to accounting periods already completed and audited.')
add_para(doc,
    'The Bidder shall also provide the name and address of the Bankers to the Bidder, and a certificate '
    'from the Chartered Accountant specifying the net worth of the Bidder and the methodology adopted '
    'for calculating such net worth.')
doc.add_paragraph()
p_yf = add_para(doc, "Yours faithfully,")
p_place = add_para(doc, "Place: {{signing_place}}")
p_date = add_para(doc, "Date: {{signing_date}}")
keep_with_next(p_yf, p_place, p_date)
add_signoff_tail(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "financial_capacity.docx"))

# ---------------------------------------------------------------------
# 11. Tender Capacity -- Airox's own project-execution track record.
#     Structurally identical every time; the actual figures (past project
#     values, work history) are your data, kept up to date by hand, not
#     something an incoming client RFP would ever contain.
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "TENDER CAPACITY")
add_para(doc, "1. Constitution or legal status of Tenderer: [Attach copy]")
add_para(doc, "Place of Registration: [Attach copy]")
add_para(doc, "Principal place of business: {{principal_place_of_business}}")
doc.add_paragraph()
add_para(doc,
    "2. Total value of Engineering construction works executed and payments received in the last five "
    "years (in Rs. Lakhs)")
table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
for yr in ["2020-21", "2021-22", "2022-23", "2023-24", "2024-25"]:
    cells = table.add_row().cells
    cells[0].text = yr
    cells[1].text = "{{value_" + yr.replace("-", "_") + "}}"
doc.add_paragraph()
add_para(doc, "Note: Attach relevant Certificate/s from Chartered Accountant.")
doc.add_paragraph()
p_yf = add_para(doc, "Yours faithfully,")
p_place = add_para(doc, "Place: {{signing_place}}")
p_date = add_para(doc, "Date: {{signing_date}}")
keep_with_next(p_yf, p_place, p_date)
add_signoff_tail(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "tender_capacity.docx"))

# ---------------------------------------------------------------------
# 12. No Deviation Certificate -- almost entirely static (a table of
#     annex references, all "NIL" every time); only the date changes.
# ---------------------------------------------------------------------
doc = base_doc()
add_title(doc, "NO DEVIATION CERTIFICATION")
add_para(doc,
    f'We, M/s. {COMPANY["company_name"]}, having its Registered Office {COMPANY["registered_office"]}, '
    'hereby certify and confirm that we have read the clauses and provisions of the RFP, amendments, '
    'addendums & clarifications issued thereafter and the stipulation of all clauses and provisions are '
    'acceptable to us, and we have not taken any deviation whatsoever to any of the clauses and '
    'provisions:')
add_para(doc, "Declaration of deviation considered by the Bidding entity, if any")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Name of the Conditions", "Reference Clause No.", "Deviation Considered"
conditions = [
    "Bid Security/EMD", "Form of tender", "Appendix I - Letter of Bid", "Annex I - Details of Bidder",
    "Annex II - Declaration for proposed technology", "Annex III - Technical Capacity",
    "Annex IV - Financial Capacity", "Annex IV(A) - CA Certificate for Financial Capacity",
    "Annex V - Statement of Legal Capacity", "Annex VI - Anti-Collusion Certificate",
    "Annex VII - Anti-Blacklisting Affidavit", "Annex VIII - Declaration of Shareholding Pattern",
    "Annex IX - Information to be Furnished by the Bidder", "Annex X - No Deviation Certificate",
    "Annex XI - Certification for not availing subsidy", "Annexure XII - Land-border procurement restriction",
    "Annexure XIII - Land-border sourcing restriction", "Appendix II - Power of Attorney",
    "Appendix III - Site Details", "Any other deviation in RFP",
]
for c in conditions:
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = c, "NIL", "NIL"
doc.add_paragraph()
add_para(doc,
    'We further confirm that we are aware that our Bid for the Project would be liable for rejection in '
    'case any material misrepresentation is made or discovered with regard to the requirements of this '
    'RFP at any stage of the Bidding Process or thereafter during the agreement period.')
doc.add_paragraph()
p_dated = add_para(doc, "Dated this {{signing_date}}")
keep_with_next(p_dated)
add_signoff_tail(doc)
ensure_unique_drawing_ids(doc)
doc.save(os.path.join(STORE, "no_deviation_certificate.docx"))

print("Templates written to", STORE)
print(os.listdir(STORE))
